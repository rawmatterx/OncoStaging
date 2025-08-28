"""
Production-Ready Modular OncoStaging Application
Simplified version with enhanced LLM prompts and core modules
"""

import streamlit as st
import logging
import json
import re
import time
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime
from dataclasses import dataclass
import tempfile
import os
import io
import base64

# Core libraries
import cv2
import numpy as np
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import pandas as pd

# AI integration
from ai_integration import MedicalAIAssistant
from config import setup_logging, PAGE_TITLE, PAGE_ICON, MAX_FILE_SIZE_BYTES, ERROR_MESSAGES
from exceptions import DocumentProcessingError, FeatureExtractionError

# Setup logging
setup_logging()
logger = logging.getLogger(__name__)


@dataclass
class OCRResult:
    """Structure for OCR processing results."""
    text: str
    confidence: float
    method: str
    processing_time: float
    page_count: int
    file_info: Dict[str, Any]


@dataclass
class ExtractionResult:
    """Structured result from extraction prompt."""
    patient_name: str = ""
    age: str = ""
    gender: str = ""
    patient_id: str = ""
    scan_date: str = ""
    cancer_type: str = ""
    tumor_location: str = ""
    tumor_size_cm: str = ""
    suv_max: str = ""
    lymph_node_involvement: Dict[str, str] = None
    distant_metastasis: Dict[str, str] = None
    clinical_impression: str = ""
    tnm_details: str = ""
    summary: str = ""
    
    def __post_init__(self):
        if self.lymph_node_involvement is None:
            self.lymph_node_involvement = {"present": "", "description": ""}
        if self.distant_metastasis is None:
            self.distant_metastasis = {"present": "", "description": ""}
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "patient_name": self.patient_name,
            "age": self.age,
            "gender": self.gender,
            "patient_id": self.patient_id,
            "scan_date": self.scan_date,
            "cancer_type": self.cancer_type,
            "tumor_location": self.tumor_location,
            "tumor_size_cm": self.tumor_size_cm,
            "suv_max": self.suv_max,
            "lymph_node_involvement": self.lymph_node_involvement,
            "distant_metastasis": self.distant_metastasis,
            "clinical_impression": self.clinical_impression,
            "tnm_details": self.tnm_details,
            "summary": self.summary
        }


@dataclass
class ClinicalRecommendation:
    """Structured clinical recommendation result."""
    ajcc_stage: str = ""
    stage_group: str = ""
    diagnostic_recommendations: List[str] = None
    treatment_recommendations: List[str] = None
    clinical_rationale: str = ""
    
    def __post_init__(self):
        if self.diagnostic_recommendations is None:
            self.diagnostic_recommendations = []
        if self.treatment_recommendations is None:
            self.treatment_recommendations = []
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "ajcc_stage": self.ajcc_stage,
            "stage_group": self.stage_group,
            "diagnostic_recommendations": self.diagnostic_recommendations,
            "treatment_recommendations": self.treatment_recommendations,
            "clinical_rationale": self.clinical_rationale
        }


class SimplifiedOCRProcessor:
    """Simplified OCR processor using Tesseract."""
    
    def __init__(self):
        self.config_options = {
            'medical_default': '--psm 6 -l eng',
            'medical_sparse': '--psm 4 -l eng',
            'medical_dense': '--psm 1 -l eng'
        }
    
    def preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Basic image preprocessing for OCR."""
        # Convert to grayscale
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image.copy()
        
        # Noise reduction
        denoised = cv2.fastNlMeansDenoising(gray)
        
        # Threshold
        _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return binary
    
    def extract_text_from_image(self, image: Image.Image) -> Tuple[str, float]:
        """Extract text from image using Tesseract."""
        try:
            # Convert to OpenCV format
            opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Preprocess
            processed = self.preprocess_image(opencv_image)
            
            # Extract text
            text = pytesseract.image_to_string(processed, config=self.config_options['medical_default'])
            
            # Basic confidence calculation
            confidence = min(1.0, len(text.strip()) / 100.0)  # Simple heuristic
            
            return text.strip(), confidence
            
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            raise DocumentProcessingError(f"OCR processing failed: {str(e)}")
    
    def process_pdf(self, pdf_file) -> OCRResult:
        """Process PDF file with OCR."""
        start_time = time.time()
        
        try:
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(pdf_file.read())
                tmp_file_path = tmp_file.name
            
            # Try standard text extraction first
            pdf_file.seek(0)
            pdf = fitz.open(stream=pdf_file.read(), filetype="pdf")
            
            standard_text = ""
            for page in pdf:
                page_text = page.get_text()
                if page_text.strip():
                    standard_text += page_text + "\n"
            
            pdf.close()
            
            # If sufficient text extracted, use it
            if len(standard_text.strip()) > 100:
                processing_time = time.time() - start_time
                return OCRResult(
                    text=standard_text,
                    confidence=0.95,
                    method="standard_extraction",
                    processing_time=processing_time,
                    page_count=len(pdf),
                    file_info={"name": pdf_file.name, "size": pdf_file.size, "type": pdf_file.type}
                )
            
            # Otherwise use OCR
            pdf = fitz.open(tmp_file_path)
            all_text = ""
            
            for page_num in range(len(pdf)):
                page = pdf.load_page(page_num)
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(img_data))
                
                text, conf = self.extract_text_from_image(image)
                all_text += text + "\n"
            
            pdf.close()
            
            processing_time = time.time() - start_time
            
            return OCRResult(
                text=all_text,
                confidence=0.7,
                method="tesseract_ocr",
                processing_time=processing_time,
                page_count=page_num + 1,
                file_info={"name": pdf_file.name, "size": pdf_file.size, "type": pdf_file.type}
            )
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise DocumentProcessingError(f"PDF processing failed: {str(e)}")
        
        finally:
            # Clean up temp file
            try:
                os.unlink(tmp_file_path)
            except:
                pass
    
    def process_docx(self, docx_file) -> OCRResult:
        """Process DOCX file with text extraction."""
        start_time = time.time()
        
        try:
            import docx
            from io import BytesIO
            
            # Read the file content
            file_content = docx_file.read()
            doc = docx.Document(BytesIO(file_content))
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = '\n'.join(text_parts)
            
            processing_time = time.time() - start_time
            
            return OCRResult(
                text=text,
                confidence=0.9,  # High confidence for direct text extraction
                method="docx_extraction",
                processing_time=processing_time,
                page_count=len(doc.paragraphs) // 30 + 1,  # Estimate pages
                file_info={"name": docx_file.name, "size": docx_file.size, "type": docx_file.type}
            )
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise DocumentProcessingError(f"DOCX processing failed: {str(e)}")
    
    def process_image(self, uploaded_file) -> OCRResult:
        """Process image file with OCR."""
        start_time = time.time()
        
        try:
            image = Image.open(uploaded_file)
            text, confidence = self.extract_text_from_image(image)
            
            processing_time = time.time() - start_time
            
            return OCRResult(
                text=text,
                confidence=confidence,
                method="tesseract_ocr",
                processing_time=processing_time,
                page_count=1,
                file_info={"name": uploaded_file.name, "size": uploaded_file.size, "type": uploaded_file.type}
            )
            
        except Exception as e:
            logger.error(f"Image processing failed: {e}")
            raise DocumentProcessingError(f"Image processing failed: {str(e)}")


class EnhancedPromptEngine:
    """Enhanced LLM prompt engine with clinical-grade prompts."""
    
    def __init__(self):
        self.ai_assistant = MedicalAIAssistant()
    
    def create_extraction_prompt(self, ocr_text: str) -> str:
        """Create comprehensive extraction prompt."""
        prompt = f"""You are a medical AI assistant that extracts structured information from PET scan reports.

Given the raw text from an OCR scan of a PET report, extract the following key data points in JSON format:
- Patient Name
- Patient Age  
- Gender
- Patient ID (if available)
- Scan Date
- Cancer Type / Diagnosis
- Tumor Location
- Tumor Size (cm)
- SUVmax (Standard Uptake Value)
- Lymph Node Involvement (Yes/No + description)
- Distant Metastasis (Yes/No + description)
- Clinical Impression
- TNM Details (if present)
- Report Summary

### Raw OCR Text:
{ocr_text}

### Important Instructions:
1. Extract only information explicitly mentioned in the text
2. Use "Not specified" for missing information
3. For numeric values, extract exact numbers when available
4. For Yes/No fields, use "Yes", "No", or "Not specified"
5. Preserve medical terminology and acronyms accurately
6. If multiple values exist for same field, use the most relevant one
7. For dates, use YYYY-MM-DD format when possible

### Output format (JSON):
{{
  "patient_name": "",
  "age": "",
  "gender": "",
  "patient_id": "",
  "scan_date": "",
  "cancer_type": "",
  "tumor_location": "",
  "tumor_size_cm": "",
  "suv_max": "",
  "lymph_node_involvement": {{
    "present": "",
    "description": ""
  }},
  "distant_metastasis": {{
    "present": "",
    "description": ""
  }},
  "clinical_impression": "",
  "tnm_details": "",
  "summary": ""
}}

JSON OUTPUT:"""
        
        return prompt
    
    def create_recommendation_prompt(self, patient_data: Dict[str, Any]) -> str:
        """Create clinical recommendation prompt."""
        # Extract data with fallbacks
        age = patient_data.get("age", "Not specified")
        gender = patient_data.get("gender", "Not specified")
        cancer_type = patient_data.get("cancer_type", "Not specified")
        tumor_size = patient_data.get("tumor_size_cm", "Not specified")
        tumor_location = patient_data.get("tumor_location", "Not specified")
        suv_max = patient_data.get("suv_max", "Not specified")
        
        # Handle complex fields
        lymph_node_data = patient_data.get("lymph_node_involvement", {})
        if isinstance(lymph_node_data, dict):
            lymph_node_desc = lymph_node_data.get("description", "Not specified")
        else:
            lymph_node_desc = str(lymph_node_data) if lymph_node_data else "Not specified"
        
        metastasis_data = patient_data.get("distant_metastasis", {})
        if isinstance(metastasis_data, dict):
            metastasis_desc = metastasis_data.get("description", "Not specified")
        else:
            metastasis_desc = str(metastasis_data) if metastasis_data else "Not specified"
        
        tnm_details = patient_data.get("tnm_details", "Not specified")
        
        prompt = f"""You are a clinical oncology assistant with access to the latest NCCN Guidelines (2025) and AJCC TNM 9th Edition.

Using the following patient cancer data, provide:
1. AJCC TNM interpretation and Stage Group
2. Recommended next steps for diagnostics or staging (based on NCCN)
3. Initial treatment options and clinical rationale

### Patient Data:
- Age: {age}
- Gender: {gender}
- Cancer Type: {cancer_type}
- Tumor Size: {tumor_size} cm
- Tumor Location: {tumor_location}
- SUVmax: {suv_max}
- Lymph Node Involvement: {lymph_node_desc}
- Distant Metastasis: {metastasis_desc}
- TNM Classification: {tnm_details}

### Clinical Guidelines to Reference:
- NCCN Clinical Practice Guidelines in Oncology (2025)
- AJCC Cancer Staging Manual, 9th Edition
- Current evidence-based treatment protocols

### Instructions:
1. Provide accurate AJCC staging interpretation
2. Recommend evidence-based diagnostic workup
3. Suggest appropriate treatment options per NCCN guidelines
4. Include clinical rationale for recommendations
5. Note any limitations due to incomplete data
6. Emphasize multidisciplinary care when appropriate

### Expected Output Format (JSON):
{{
  "ajcc_stage": "",
  "stage_group": "",
  "diagnostic_recommendations": [],
  "treatment_recommendations": [],
  "clinical_rationale": "",
  "data_limitations": "",
  "multidisciplinary_care": "",
  "follow_up_recommendations": []
}}

JSON OUTPUT:"""
        
        return prompt
    
    def extract_structured_data(self, ocr_text: str) -> ExtractionResult:
        """Extract structured data using enhanced prompts."""
        try:
            prompt = self.create_extraction_prompt(ocr_text)
            
            response = self.ai_assistant.client.query_model(
                prompt=prompt,
                model="deepseek-chat",
                temperature=0.1,
                max_tokens=2048
            )
            
            # Parse JSON response
            json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                data = json.loads(json_str)
            else:
                raise ValueError("No JSON found in response")
            
            # Create ExtractionResult
            result = ExtractionResult(
                patient_name=data.get("patient_name", ""),
                age=data.get("age", ""),
                gender=data.get("gender", ""),
                patient_id=data.get("patient_id", ""),
                scan_date=data.get("scan_date", ""),
                cancer_type=data.get("cancer_type", ""),
                tumor_location=data.get("tumor_location", ""),
                tumor_size_cm=data.get("tumor_size_cm", ""),
                suv_max=data.get("suv_max", ""),
                lymph_node_involvement=data.get("lymph_node_involvement", {}),
                distant_metastasis=data.get("distant_metastasis", {}),
                clinical_impression=data.get("clinical_impression", ""),
                tnm_details=data.get("tnm_details", ""),
                summary=data.get("summary", "")
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Enhanced extraction failed: {e}")
            raise FeatureExtractionError(f"Extraction failed: {str(e)}")
    
    def generate_clinical_recommendations(self, patient_data: Dict[str, Any]) -> ClinicalRecommendation:
        """Generate clinical recommendations."""
        try:
            prompt = self.create_recommendation_prompt(patient_data)
            
            response = self.ai_assistant.client.query_model(
                prompt=prompt,
                model="deepseek-chat",
                temperature=0.2,
                max_tokens=2048
            )
            
            # Parse JSON response
            json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                data = json.loads(json_str)
            else:
                raise ValueError("No JSON found in response")
            
            # Create ClinicalRecommendation
            result = ClinicalRecommendation(
                ajcc_stage=data.get("ajcc_stage", ""),
                stage_group=data.get("stage_group", ""),
                diagnostic_recommendations=data.get("diagnostic_recommendations", []),
                treatment_recommendations=data.get("treatment_recommendations", []),
                clinical_rationale=data.get("clinical_rationale", "")
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Clinical recommendation failed: {e}")
            raise FeatureExtractionError(f"Recommendation generation failed: {str(e)}")


class ProductionOncoStagingApp:
    """Production-ready OncoStaging application."""
    
    def __init__(self):
        self.ocr_processor = SimplifiedOCRProcessor()
        self.prompt_engine = EnhancedPromptEngine()
        self._initialize_session_state()
    
    def _initialize_session_state(self):
        """Initialize session state."""
        if 'prod_app_state' not in st.session_state:
            st.session_state.prod_app_state = {
                'ocr_result': None,
                'extraction_result': None,
                'clinical_recommendations': None,
                'current_step': 1
            }
    
    def run(self):
        """Run the production application."""
        # Page configuration
        st.set_page_config(
            page_title="Production OncoStaging",
            page_icon="🏥",
            layout="wide",
            initial_sidebar_state="expanded"
        )
        
        # Header
        st.title("🏥 Production OncoStaging - Clinical AI Platform")
        st.markdown("**Advanced PET scan analysis with enhanced LLM prompts and clinical guidelines**")
        
        # Navigation
        self._render_navigation()
        
        # Main content
        current_step = st.session_state.prod_app_state['current_step']
        
        if current_step == 1:
            self._step1_ocr_processing()
        elif current_step == 2:
            self._step2_data_extraction()
        elif current_step == 3:
            self._step3_clinical_analysis()
        elif current_step == 4:
            self._step4_final_report()
    
    def _render_navigation(self):
        """Render navigation sidebar."""
        with st.sidebar:
            st.header("🎯 Process Flow")
            
            steps = [
                ("1️⃣ OCR Processing", 1),
                ("2️⃣ Data Extraction", 2),
                ("3️⃣ Clinical Analysis", 3),
                ("4️⃣ Final Report", 4)
            ]
            
            current = st.session_state.prod_app_state['current_step']
            
            for step_name, step_num in steps:
                # Status indicator
                if step_num < current:
                    status = "✅"
                elif step_num == current:
                    status = "🔄"
                else:
                    status = "⏳"
                
                # Navigation button
                if st.button(f"{status} {step_name}", key=f"nav_{step_num}"):
                    if step_num <= current or st.session_state.prod_app_state.get('ocr_result'):
                        st.session_state.prod_app_state['current_step'] = step_num
                        st.rerun()
            
            # Data status
            st.markdown("---")
            st.subheader("📊 Status")
            
            state = st.session_state.prod_app_state
            
            if state['ocr_result']:
                st.success("📄 OCR Complete")
            else:
                st.info("📄 OCR Pending")
            
            if state['extraction_result']:
                st.success("🔬 Extraction Complete")
            else:
                st.info("🔬 Extraction Pending")
            
            if state['clinical_recommendations']:
                st.success("🏥 Analysis Complete")
            else:
                st.info("🏥 Analysis Pending")
            
            # Reset
            st.markdown("---")
            if st.button("🔄 Reset"):
                self._reset_application()
                st.rerun()
    
    def _step1_ocr_processing(self):
        """Step 1: OCR Processing."""
        st.header("📄 Step 1: Document Upload & OCR")
        
        # File upload
        uploaded_file = st.file_uploader(
            "Upload your PET scan report",
            type=['pdf', 'docx', 'png', 'jpg', 'jpeg', 'tiff', 'bmp'],
            help="Supports PDF, DOCX documents, and image files"
        )
        
        if uploaded_file is not None:
            # File validation
            if uploaded_file.size > MAX_FILE_SIZE_BYTES:
                st.error(f"❌ File too large. Maximum size: {MAX_FILE_SIZE_BYTES/(1024*1024):.1f}MB")
                return
            
            # Display file info
            st.info(f"**File:** {uploaded_file.name} | **Size:** {uploaded_file.size:,} bytes")
            
            # Process button
            if st.button("🔍 Process Document", type="primary"):
                with st.spinner("Processing document with OCR..."):
                    try:
                        # Process based on file type
                        file_ext = uploaded_file.name.split('.')[-1].lower()
                        
                        if file_ext == 'pdf':
                            ocr_result = self.ocr_processor.process_pdf(uploaded_file)
                        elif file_ext == 'docx':
                            ocr_result = self.ocr_processor.process_docx(uploaded_file)
                        else:
                            ocr_result = self.ocr_processor.process_image(uploaded_file)
                        
                        # Store result
                        st.session_state.prod_app_state['ocr_result'] = ocr_result
                        
                        # Display results
                        self._display_ocr_results(ocr_result)
                        
                        # Auto-advance
                        if st.button("➡️ Proceed to Data Extraction", type="primary"):
                            st.session_state.prod_app_state['current_step'] = 2
                            st.rerun()
                        
                    except Exception as e:
                        st.error(f"❌ OCR processing failed: {str(e)}")
        
        # Show sample
        elif not st.session_state.prod_app_state.get('ocr_result'):
            self._show_sample_data()
    
    def _step2_data_extraction(self):
        """Step 2: Medical Data Extraction."""
        st.header("🔬 Step 2: Advanced Medical Data Extraction")
        
        ocr_result = st.session_state.prod_app_state.get('ocr_result')
        
        if not ocr_result:
            st.warning("⚠️ Please complete OCR processing first.")
            if st.button("← Back to OCR"):
                st.session_state.prod_app_state['current_step'] = 1
                st.rerun()
            return
        
        # Display OCR summary
        st.info(f"**Source:** {ocr_result.file_info['name']} | **Method:** {ocr_result.method} | **Confidence:** {ocr_result.confidence:.1%}")
        
        # Extraction button
        if st.button("🔍 Extract Medical Data", type="primary"):
            with st.spinner("Extracting structured medical data with enhanced LLM prompts..."):
                try:
                    extraction_result = self.prompt_engine.extract_structured_data(ocr_result.text)
                    
                    # Store result
                    st.session_state.prod_app_state['extraction_result'] = extraction_result
                    
                    # Display results
                    self._display_extraction_results(extraction_result)
                    
                    # Auto-advance
                    if st.button("➡️ Proceed to Clinical Analysis", type="primary"):
                        st.session_state.prod_app_state['current_step'] = 3
                        st.rerun()
                    
                except Exception as e:
                    st.error(f"❌ Extraction failed: {str(e)}")
        
        # Show previous results
        if st.session_state.prod_app_state.get('extraction_result'):
            st.markdown("---")
            st.subheader("📋 Previous Extraction Results")
            self._display_extraction_results(st.session_state.prod_app_state['extraction_result'])
    
    def _step3_clinical_analysis(self):
        """Step 3: Clinical Analysis & Recommendations."""
        st.header("🏥 Step 3: Clinical Analysis & Recommendations")
        
        extraction_result = st.session_state.prod_app_state.get('extraction_result')
        
        if not extraction_result:
            st.warning("⚠️ Please complete data extraction first.")
            if st.button("← Back to Extraction"):
                st.session_state.prod_app_state['current_step'] = 2
                st.rerun()
            return
        
        # Display extraction summary
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Cancer Type", extraction_result.cancer_type or "Not specified")
        with col2:
            st.metric("Tumor Size", extraction_result.tumor_size_cm or "Not specified")
        with col3:
            st.metric("SUV Max", extraction_result.suv_max or "Not specified")
        
        # Generate recommendations
        if st.button("🔍 Generate Clinical Recommendations", type="primary"):
            with st.spinner("Generating clinical recommendations with NCCN guidelines..."):
                try:
                    # Prepare patient data
                    patient_data = extraction_result.to_dict()
                    
                    # Generate recommendations
                    clinical_recommendations = self.prompt_engine.generate_clinical_recommendations(patient_data)
                    
                    # Store result
                    st.session_state.prod_app_state['clinical_recommendations'] = clinical_recommendations
                    
                    # Display results
                    self._display_clinical_recommendations(clinical_recommendations)
                    
                    # Auto-advance
                    if st.button("➡️ Generate Final Report", type="primary"):
                        st.session_state.prod_app_state['current_step'] = 4
                        st.rerun()
                    
                except Exception as e:
                    st.error(f"❌ Clinical analysis failed: {str(e)}")
        
        # Show previous results
        if st.session_state.prod_app_state.get('clinical_recommendations'):
            st.markdown("---")
            st.subheader("📋 Previous Clinical Analysis")
            self._display_clinical_recommendations(st.session_state.prod_app_state['clinical_recommendations'])
    
    def _step4_final_report(self):
        """Step 4: Final Report Generation."""
        st.header("📄 Step 4: Final Report & Export")
        
        # Check data availability
        state = st.session_state.prod_app_state
        
        if not all([state.get('ocr_result'), state.get('extraction_result'), state.get('clinical_recommendations')]):
            st.warning("⚠️ Please complete all previous steps first.")
            return
        
        # Generate comprehensive report
        report_data = self._generate_comprehensive_report()
        
        # Display executive summary
        st.subheader("📋 Executive Summary")
        
        extraction = state['extraction_result']
        clinical = state['clinical_recommendations']
        
        summary = f"""
**PATIENT INFORMATION**
- Age: {extraction.age or 'Not specified'}
- Gender: {extraction.gender or 'Not specified'}
- Cancer Type: {extraction.cancer_type or 'Not specified'}
- Tumor Location: {extraction.tumor_location or 'Not specified'}

**IMAGING FINDINGS**
- Tumor Size: {extraction.tumor_size_cm or 'Not specified'}
- SUV Max: {extraction.suv_max or 'Not specified'}
- TNM Details: {extraction.tnm_details or 'Not specified'}

**CLINICAL ASSESSMENT**
- AJCC Stage: {clinical.ajcc_stage}
- Stage Group: {clinical.stage_group}

**CLINICAL IMPRESSION**
{extraction.clinical_impression or 'Not specified'}

**ANALYSIS CONFIDENCE**
- OCR Quality: {state['ocr_result'].confidence:.1%}
- Processing Method: {state['ocr_result'].method.replace('_', ' ').title()}

*This analysis was performed using AI-assisted extraction and NCCN/AJCC clinical guidelines. All recommendations should be reviewed by qualified oncology professionals.*
"""
        
        st.text_area("Executive Summary", summary, height=400)
        
        # Display detailed sections
        tabs = st.tabs(["🔬 Extraction Data", "🏥 Clinical Recommendations", "📊 Processing Details"])
        
        with tabs[0]:
            st.json(extraction.to_dict())
        
        with tabs[1]:
            st.json(clinical.to_dict())
        
        with tabs[2]:
            st.json({
                "ocr_processing": {
                    "method": state['ocr_result'].method,
                    "confidence": state['ocr_result'].confidence,
                    "processing_time": state['ocr_result'].processing_time,
                    "page_count": state['ocr_result'].page_count
                },
                "file_info": state['ocr_result'].file_info
            })
        
        # Export options
        self._render_export_options(report_data, summary)
    
    def _display_ocr_results(self, result: OCRResult):
        """Display OCR processing results."""
        st.success("✅ OCR processing completed!")
        
        # Metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Confidence", f"{result.confidence:.1%}")
        with col2:
            st.metric("Processing Time", f"{result.processing_time:.2f}s")
        with col3:
            st.metric("Method", result.method.replace("_", " ").title())
        with col4:
            st.metric("Pages", result.page_count)
        
        # Preview text
        with st.expander("📝 Preview Extracted Text", expanded=True):
            text_length = len(result.text)
            if text_length > 5000:
                st.text_area("Extracted Text (First 5000 characters)", result.text[:5000] + "...", height=300)
                st.info(f"Full text contains {text_length:,} characters.")
            else:
                st.text_area("Extracted Text", result.text, height=300)
    
    def _display_extraction_results(self, result: ExtractionResult):
        """Display extraction results."""
        st.success("✅ Medical data extraction completed!")
        
        # Key findings
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Cancer Type", result.cancer_type or "Not specified")
        with col2:
            st.metric("Tumor Size", result.tumor_size_cm or "Not specified")
        with col3:
            st.metric("SUV Max", result.suv_max or "Not specified")
        
        # Detailed results
        with st.expander("🔍 Detailed Extraction Results", expanded=True):
            
            # Patient Info
            st.subheader("👤 Patient Information")
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"**Name:** {result.patient_name or 'Not specified'}")
                st.write(f"**Age:** {result.age or 'Not specified'}")
            with col2:
                st.write(f"**Gender:** {result.gender or 'Not specified'}")
                st.write(f"**Patient ID:** {result.patient_id or 'Not specified'}")
            
            # Cancer Info
            st.subheader("🎯 Cancer Information")
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"**Type:** {result.cancer_type or 'Not specified'}")
                st.write(f"**Location:** {result.tumor_location or 'Not specified'}")
            with col2:
                st.write(f"**Size:** {result.tumor_size_cm or 'Not specified'}")
                st.write(f"**SUV Max:** {result.suv_max or 'Not specified'}")
            
            # Clinical Findings
            st.subheader("🩺 Clinical Findings")
            st.write(f"**TNM Details:** {result.tnm_details or 'Not specified'}")
            st.write(f"**Clinical Impression:** {result.clinical_impression or 'Not specified'}")
            
            # Lymph Nodes & Metastasis
            if result.lymph_node_involvement:
                st.write(f"**Lymph Node Involvement:** {result.lymph_node_involvement.get('present', 'Not specified')}")
                if result.lymph_node_involvement.get('description'):
                    st.write(f"**Description:** {result.lymph_node_involvement['description']}")
            
            if result.distant_metastasis:
                st.write(f"**Distant Metastasis:** {result.distant_metastasis.get('present', 'Not specified')}")
                if result.distant_metastasis.get('description'):
                    st.write(f"**Description:** {result.distant_metastasis['description']}")
    
    def _display_clinical_recommendations(self, result: ClinicalRecommendation):
        """Display clinical recommendations."""
        st.success("✅ Clinical recommendations generated!")
        
        # AJCC Staging
        col1, col2 = st.columns(2)
        with col1:
            st.metric("AJCC Stage", result.ajcc_stage or "Not specified")
        with col2:
            st.metric("Stage Group", result.stage_group or "Not specified")
        
        # Recommendations
        if result.diagnostic_recommendations:
            st.subheader("🔍 Diagnostic Recommendations")
            for rec in result.diagnostic_recommendations:
                st.write(f"• {rec}")
        
        if result.treatment_recommendations:
            st.subheader("💊 Treatment Recommendations")
            for rec in result.treatment_recommendations:
                st.write(f"• {rec}")
        
        # Clinical Rationale
        if result.clinical_rationale:
            st.subheader("🧠 Clinical Rationale")
            st.write(result.clinical_rationale)
    
    def _generate_comprehensive_report(self) -> Dict[str, Any]:
        """Generate comprehensive report."""
        state = st.session_state.prod_app_state
        
        return {
            "report_metadata": {
                "generation_timestamp": datetime.now().isoformat(),
                "report_type": "Production OncoStaging Analysis",
                "version": "1.0"
            },
            "patient_data": state['extraction_result'].to_dict(),
            "clinical_analysis": state['clinical_recommendations'].to_dict(),
            "processing_info": {
                "ocr_method": state['ocr_result'].method,
                "ocr_confidence": state['ocr_result'].confidence,
                "processing_time": state['ocr_result'].processing_time
            }
        }
    
    def _render_export_options(self, report_data: Dict[str, Any], summary: str):
        """Render export options."""
        st.subheader("📥 Export Options")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # JSON Export
            json_data = json.dumps(report_data, indent=2, default=str)
            st.download_button(
                "📄 Download JSON Report",
                data=json_data,
                file_name=f"oncostaging_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )
        
        with col2:
            # Summary Export
            st.download_button(
                "📝 Download Summary",
                data=summary,
                file_name=f"oncostaging_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain"
            )
        
        with col3:
            if st.button("📧 Email Report"):
                st.info("Email functionality would be implemented here")
    
    def _show_sample_data(self):
        """Show sample data for demonstration."""
        with st.expander("📋 View Sample Analysis", expanded=False):
            st.info("**Sample PET/CT Report Text:**")
            sample_text = """
Patient: John Doe, Age: 65, Male
Study Date: 2024-01-15
Indication: Lung cancer staging

FINDINGS:
There is a hypermetabolic mass in the right upper lobe measuring approximately 3.2 x 2.8 cm with SUVmax of 8.5.
Multiple enlarged mediastinal lymph nodes with increased FDG uptake, largest measuring 1.5 cm with SUVmax of 4.2.
No evidence of distant metastatic disease.

IMPRESSION:
1. Right upper lobe mass consistent with primary lung cancer (T2N2M0, Stage IIIA)
2. Mediastinal lymphadenopathy consistent with nodal metastases
3. No distant metastases identified

Recommend tissue confirmation and multidisciplinary evaluation.
"""
            st.text_area("Sample Report", sample_text, height=200)
            
            if st.button("🧪 Use Sample Data"):
                # Create mock OCR result
                mock_ocr = OCRResult(
                    text=sample_text,
                    confidence=0.95,
                    method="sample_data",
                    processing_time=0.1,
                    page_count=1,
                    file_info={"name": "sample_report.txt", "size": len(sample_text), "type": "text/plain"}
                )
                
                st.session_state.prod_app_state['ocr_result'] = mock_ocr
                st.success("✅ Sample data loaded!")
                st.info("You can now proceed to Step 2: Data Extraction")
    
    def _reset_application(self):
        """Reset application state."""
        st.session_state.prod_app_state = {
            'ocr_result': None,
            'extraction_result': None,
            'clinical_recommendations': None,
            'current_step': 1
        }


def main():
    """Main entry point."""
    app = ProductionOncoStagingApp()
    app.run()


if __name__ == "__main__":
    main()
