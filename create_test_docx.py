from docx import Document

def create_test_docx(file_path):
    doc = Document()
    doc.add_heading('Test Document for OncoStaging', 0)
    
    doc.add_paragraph('Patient Information:')
    doc.add_paragraph('Name: John Doe')
    doc.add_paragraph('Age: 45')
    doc.add_paragraph('Gender: Male')
    
    doc.add_heading('Medical History', level=1)
    doc.add_paragraph('- Hypertension')
    doc.add_paragraph('- Type 2 Diabetes')
    
    doc.add_heading('Current Symptoms', level=1)
    doc.add_paragraph('- Fatigue')
    doc.add_paragraph('- Weight loss')
    doc.add_paragraph('- Abdominal pain')
    
    doc.save(file_path)
    print(f"Created test document: {file_path}")

if __name__ == "__main__":
    create_test_docx("test_document_proper.docx")
