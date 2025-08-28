"""
Document processing module for OncoStaging application.
Handles file validation, text extraction, and document processing.
"""

import os
import hashlib
import logging
from typing import Optional, Dict, Any
from pathlib import Path
import json
import time

import fitz  # PyMuPDF
import docx
import streamlit as st

from config import (
    MAX_FILE_SIZE_BYTES, ALLOWED_FILE_TYPES, ALLOWED_MIME_TYPES,
    CACHE_ENABLED, CACHE_DIR, CACHE_TTL_HOURS, MIN_TEXT_LENGTH,
    ERROR_MESSAGES
)
from exceptions import (
    DocumentProcessingError, FileValidationError,
    OncoStagingError
)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document validation, processing, and text extraction."""
    
    def __init__(self):
        """Initialize the document processor."""
        self.cache_dir = Path(CACHE_DIR)
        if CACHE_ENABLED:
            self.cache_dir.mkdir(exist_ok=True)
    
    def validate_file(self, file) -> None:
        """
        Validate uploaded file for type, size, and content.
        
        Args:
            file: Streamlit UploadedFile object
            
        Raises:
            FileValidationError: If file validation fails
        """
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > MAX_FILE_SIZE_BYTES:
            max_mb = MAX_FILE_SIZE_BYTES / (1024 * 1024)
            raise FileValidationError(
                ERROR_MESSAGES["file_too_large"].format(max_size=max_mb)
            )
        
        # Check file extension
        file_ext = file.name.split('.')[-1].lower()
        if file_ext not in ALLOWED_FILE_TYPES:
            raise FileValidationError(ERROR_MESSAGES["invalid_file_type"])
        
        # Check MIME type if available
        if hasattr(file, 'type') and file.type not in ALLOWED_MIME_TYPES:
            raise FileValidationError(ERROR_MESSAGES["invalid_file_type"])
        
        # Check if file is not empty
        if file_size == 0:
            raise FileValidationError("ফাইলটি খালি। অনুগ্রহ করে একটি বৈধ ডকুমেন্ট আপলোড করুন।")
        
        logger.info(f"File validation successful: {file.name} ({file_size} bytes)")
    
    def get_file_hash(self, file) -> str:
        """
        Calculate SHA256 hash of the file.
        
        Args:
            file: File object
            
        Returns:
            str: SHA256 hash of the file
        """
        file.seek(0)
        file_hash = hashlib.sha256(file.read()).hexdigest()
        file.seek(0)
        return file_hash
    
    def get_cache_path(self, file_hash: str) -> Path:
        """Get cache file path for given file hash."""
        return self.cache_dir / f"{file_hash}.json"
    
    def load_from_cache(self, file_hash: str) -> Optional[Dict[str, Any]]:
        """
        Load processed document from cache if available and not expired.
        
        Args:
            file_hash: SHA256 hash of the file
            
        Returns:
            Cached data if valid, None otherwise
        """
        if not CACHE_ENABLED:
            return None
        
        cache_path = self.get_cache_path(file_hash)
        
        if not cache_path.exists():
            return None
        
        try:
            with open(cache_path, 'r', encoding='utf-8') as f:
                cached_data = json.load(f)
            
            # Check if cache is expired
            cache_time = cached_data.get('timestamp', 0)
            current_time = time.time()
            cache_age_hours = (current_time - cache_time) / 3600
            
            if cache_age_hours > CACHE_TTL_HOURS:
                logger.info(f"Cache expired for hash {file_hash}")
                cache_path.unlink()  # Delete expired cache
                return None
            
            logger.info(f"Loading from cache: {file_hash}")
            return cached_data.get('data')
            
        except Exception as e:
            logger.error(f"Error loading cache: {e}")
            return None
    
    def save_to_cache(self, file_hash: str, data: Dict[str, Any]) -> None:
        """
        Save processed document to cache.
        
        Args:
            file_hash: SHA256 hash of the file
            data: Data to cache
        """
        if not CACHE_ENABLED:
            return
        
        cache_path = self.get_cache_path(file_hash)
        
        try:
            cache_data = {
                'timestamp': time.time(),
                'data': data
            }
            
            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump(cache_data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"Saved to cache: {file_hash}")
            
        except Exception as e:
            logger.error(f"Error saving to cache: {e}")
    
    def extract_text_from_pdf(self, file) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file: PDF file object
            
        Returns:
            Extracted text
            
        Raises:
            DocumentProcessingError: If PDF processing fails
        """
        try:
            pdf = fitz.open(stream=file.read(), filetype="pdf")
            text_parts = []
            
            for page_num, page in enumerate(pdf):
                try:
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
            
            pdf.close()
            
            if not text_parts:
                raise DocumentProcessingError("PDF থেকে কোন টেক্সট এক্সট্র্যাক্ট করা যায়নি।")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise DocumentProcessingError(f"PDF প্রসেসিং এ সমস্যা: {str(e)}")
    
    def extract_text_from_docx(self, file) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file: DOCX file object
            
        Returns:
            Extracted text
            
        Raises:
            DocumentProcessingError: If DOCX processing fails
        """
        try:
            doc = docx.Document(file)
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            if not text_parts:
                raise DocumentProcessingError("DOCX থেকে কোন টেক্সট এক্সট্র্যাক্ট করা যায়নি।")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise DocumentProcessingError(f"DOCX প্রসেসিং এ সমস্যা: {str(e)}")
    
    @st.cache_data(ttl=3600)
    def process_document(_self, file) -> Dict[str, Any]:
        """
        Process uploaded document and extract text.
        
        Args:
            file: Uploaded file object
            
        Returns:
            Dictionary with extracted text and metadata
            
        Raises:
            OncoStagingError: If processing fails
        """
        try:
            # Validate file
            _self.validate_file(file)
            
            # Get file hash for caching
            file_hash = _self.get_file_hash(file)
            
            # Check cache
            cached_data = _self.load_from_cache(file_hash)
            if cached_data:
                return cached_data
            
            # Extract text based on file type
            file_ext = file.name.split('.')[-1].lower()
            
            if file_ext == 'pdf':
                text = _self.extract_text_from_pdf(file)
            elif file_ext == 'docx':
                text = _self.extract_text_from_docx(file)
            else:
                raise FileValidationError(ERROR_MESSAGES["invalid_file_type"])
            
            # Validate extracted text
            if len(text.strip()) < MIN_TEXT_LENGTH:
                raise DocumentProcessingError(
                    "এক্সট্র্যাক্ট করা টেক্সট খুব ছোট। অনুগ্রহ করে একটি সম্পূর্ণ মেডিকেল রিপোর্ট আপলোড করুন।"
                )
            
            # Prepare result
            result = {
                'text': text,
                'file_name': file.name,
                'file_type': file_ext,
                'file_size': file.size if hasattr(file, 'size') else 0,
                'text_length': len(text),
                'processing_time': time.time()
            }
            
            # Save to cache
            _self.save_to_cache(file_hash, result)
            
            logger.info(f"Document processed successfully: {file.name}")
            return result
            
        except OncoStagingError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error in document processing: {e}")
            raise DocumentProcessingError(ERROR_MESSAGES["processing_error"])
    
    def clear_cache(self) -> None:
        """Clear all cached files."""
        if not CACHE_ENABLED:
            return
        
        try:
            for cache_file in self.cache_dir.glob("*.json"):
                cache_file.unlink()
            logger.info("Cache cleared successfully")
        except Exception as e:
            logger.error(f"Error clearing cache: {e}")
    
    def get_cache_stats(self) -> Dict[str, Any]:
        """Get cache statistics."""
        if not CACHE_ENABLED:
            return {"enabled": False}
        
        cache_files = list(self.cache_dir.glob("*.json"))
        total_size = sum(f.stat().st_size for f in cache_files)
        
        return {
            "enabled": True,
            "files_count": len(cache_files),
            "total_size_mb": round(total_size / (1024 * 1024), 2),
            "cache_dir": str(self.cache_dir)
        }
